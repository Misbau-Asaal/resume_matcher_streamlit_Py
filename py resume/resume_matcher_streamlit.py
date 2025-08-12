"""
Resume vs Job Description Matcher
Streamlit app that accepts resume and job description files (txt/pdf/docx), parses text,
extracts keywords, computes a match score, and gives actionable feedback.

Requirements (pip install):
  pip install streamlit pdfplumber python-docx spacy rapidfuzz scikit-learn fpdf
  python -m spacy download en_core_web_sm

Run:
  streamlit run resume_matcher_streamlit.py

Notes:
- The app uses spaCy for tokenization/lemmatization and RapidFuzz for fuzzy matching.
- A small builtin skills list is provided; you can replace it with a larger taxonomy (O*NET, GitHub skills CSV).

"""

from io import BytesIO
import streamlit as st
import pdfplumber
import docx
import re
from rapidfuzz import process, fuzz
import spacy
from typing import List, Set, Tuple, Dict
from fpdf import FPDF

# Load spaCy model once
nlp = spacy.load("en_core_web_sm")

# --- Simple skills taxonomy (extend or load from external CSV) ---
BUILTIN_SKILLS = {
    'python', 'java', 'c++', 'c#', 'javascript', 'typescript', 'react', 'angular', 'vue',
    'node.js', 'node', 'express', 'django', 'flask', 'fastapi', 'sql', 'postgresql', 'mysql', 'mongodb',
    'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'ci/cd', 'jenkins', 'git', 'rest api', 'graphql',
    'html', 'css', 'sass', 'linux', 'bash', 'powershell', 'tensorflow', 'pytorch', 'nlp', 'computer vision',
    'agile', 'scrum', 'kanban', 'jira', 'trello', 'unit testing', 'pytest', 'jest', 'selenium', 'redis',
    'microservices', 'oauth', 'ldap', 'oauth2', 'orm', 'hibernate', 'spring', 'dotnet', 'dotnet core'
}

# Normalization helpers
WORD_CHAR_RE = re.compile(r"[^a-zA-Z0-9+#\- ]")


def extract_text_from_pdf(file_bytes: BytesIO) -> str:
    file_bytes.seek(0)
    text_parts = []
    with pdfplumber.open(file_bytes) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: BytesIO) -> str:
    file_bytes.seek(0)
    doc = docx.Document(file_bytes)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs)


def extract_text_from_txt(file_bytes: BytesIO) -> str:
    try:
        return file_bytes.getvalue().decode('utf-8')
    except Exception:
        return file_bytes.getvalue().decode('latin-1')


def load_file_to_text(uploaded_file) -> Tuple[str, str]:
    """Return (text, filename)"""
    if uploaded_file is None:
        return "", ""
    fname = uploaded_file.name
    fb = BytesIO(uploaded_file.read())
    if fname.lower().endswith('.pdf'):
        text = extract_text_from_pdf(fb)
    elif fname.lower().endswith('.docx'):
        text = extract_text_from_docx(fb)
    elif fname.lower().endswith('.txt'):
        text = extract_text_from_txt(fb)
    else:
        # Attempt to read as text
        text = extract_text_from_txt(fb)
    return text, fname


def normalize_text_for_skills(text: str) -> str:
    # Lowercase, remove weird chars but keep +, #, - for things like C++, C#, .NET
    cleaned = WORD_CHAR_RE.sub(' ', text)
    return cleaned.lower()


def lemmatize_tokens(text: str, nlp_model) -> List[str]:
    doc = nlp_model(text)
    toks = [token.lemma_.lower()
            for token in doc if token.is_alpha and not token.is_stop]
    return toks


def extract_candidate_terms(text: str, nlp_model) -> Set[str]:
    """Extract nouns, proper nouns, noun chunks, and named entities as candidate keywords."""
    doc = nlp_model(text)
    terms = set()
    for chunk in doc.noun_chunks:
        chunk_text = chunk.text.strip().lower()
        if len(chunk_text) > 1:
            terms.add(chunk_text)
    for ent in doc.ents:
        ent_text = ent.text.strip().lower()
        if len(ent_text) > 1:
            terms.add(ent_text)
    # also add single noun/proper-noun tokens
    for token in doc:
        if token.pos_ in ("NOUN", "PROPN") and not token.is_stop and token.is_alpha:
            terms.add(token.lemma_.lower())
    # normalize (remove trailing stopwords, short words)
    final = {re.sub(r'[^a-z0-9+#\- ]', '', t).strip() for t in terms}
    final = {t for t in final if len(t) > 1}
    return final


def match_skills_against_text(skills: Set[str], text: str, threshold: int = 85) -> Tuple[Set[str], Dict[str, Tuple[str, int]]]:
    """Return matched skills and a dict of skill->(best_match_in_text, score) for fuzzy matches."""
    text_norm = normalize_text_for_skills(text)
    words = list(set([w.strip()
                 for w in re.split(r'\s+', text_norm) if w.strip()]))
    matched = set()
    match_details = {}
    for skill in skills:
        # try exact
        skill_norm = skill.lower()
        if skill_norm in text_norm:
            matched.add(skill)
            match_details[skill] = (skill, 100)
            continue
        # fuzzy match against the text words and phrases
        choice, score, _ = process.extractOne(
            skill_norm, words, scorer=fuzz.token_sort_ratio)
        if score >= threshold:
            matched.add(skill)
            match_details[skill] = (choice, int(score))
    return matched, match_details


def build_job_keywords(job_text: str, builtin_skills: Set[str], nlp_model, fuzzy_threshold: int = 82) -> Tuple[Set[str], Dict[str, int]]:
    """Return a set of keywords for the job description and a small frequency map."""
    text_norm = normalize_text_for_skills(job_text)
    # 1) Explicit skills found from builtin list
    matched_builtin, details = match_skills_against_text(
        builtin_skills, job_text, threshold=fuzzy_threshold)

    # 2) Candidate noun phrases and entities
    candidate_terms = extract_candidate_terms(job_text, nlp_model)

    # Filter candidate terms: remove stopwords and very generic words
    generic_blacklist = {'experience', 'team', 'work', 'years',
                         'role', 'skills', 'knowledge', 'ability', 'responsibility'}
    filtered = {
        t for t in candidate_terms if t not in generic_blacklist and len(t) <= 40}

    # Combine and create frequency map
    final_keywords = set(matched_builtin)
    freq = {}
    for k in matched_builtin:
        freq[k] = freq.get(k, 0) + 1
    for term in filtered:
        # only keep multi-word terms or those with letters
        if re.search('[a-z]', term):
            final_keywords.add(term)
            freq[term] = freq.get(term, 0) + 1

    # Heuristic: prefer shorter, common skill tokens too (split multi-words)
    words = set()
    for term in final_keywords.copy():
        for w in term.split():
            if len(w) > 1:
                words.add(w)
    for w in words:
        final_keywords.add(w)
        freq[w] = freq.get(w, 0) + 0

    return final_keywords, freq


def compute_match(resume_text: str, job_keywords: Set[str], fuzzy_threshold: int = 82) -> Tuple[float, List[str], List[str], Dict[str, int]]:
    # Match builtin keywords fuzzily
    resume_norm = normalize_text_for_skills(resume_text)
    resume_terms = set(re.split(r'\s+', resume_norm))

    matched = set()
    match_scores = {}
    for kw in job_keywords:
        kw_norm = kw.lower()
        if kw_norm in resume_norm:
            matched.add(kw)
            match_scores[kw] = 100
            continue
        # fuzzy match against resume terms
        choice, score, _ = process.extractOne(
            kw_norm, list(resume_terms), scorer=fuzz.token_sort_ratio)
        if score >= fuzzy_threshold:
            matched.add(kw)
            match_scores[kw] = int(score)

    missing = sorted(list(set(job_keywords) - matched))
    extra = sorted(list({t for t in resume_terms if t and len(
        t) > 1} - {k.lower() for k in job_keywords}))[:200]

    # percent match = matched / total job keywords
    total_keywords = len(job_keywords) if job_keywords else 1
    score_pct = int(round(len(matched) / total_keywords * 100))
    return score_pct, sorted(list(matched)), missing, extra[:100]


def generate_suggestions(missing_keywords: List[str], extra_keywords: List[str]) -> List[str]:
    suggestions = []
    if missing_keywords:
        suggestions.append("Add or emphasize the following skills/terms in your resume: " +
                           ', '.join(missing_keywords[:8]) + ("..." if len(missing_keywords) > 8 else ""))
        suggestions.append(
            "Show concrete projects or bullets that demonstrate these skills (e.g., 'Built a REST API using Flask and MongoDB for ...').")
    else:
        suggestions.append(
            "Good news: your resume already contains most of the job keywords.")
    if extra_keywords:
        suggestions.append("Consider removing or de-emphasizing unrelated skills: " +
                           ', '.join(extra_keywords[:8]) + ("..." if len(extra_keywords) > 8 else ""))
    suggestions.append("If you lack a required skill, consider short projects, online courses, or a one-line note under skills indicating familiarity level (e.g., 'Familiar with MongoDB - used in side project').")
    return suggestions


def export_report_pdf(match_score, matched, missing, extra, suggestions):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=14)
    pdf.cell(200, 10, txt=f"Match Score: {match_score:.2f}%", ln=True)

    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 8, f"Matched Keywords: {', '.join(matched) or 'None'}")
    pdf.multi_cell(0, 8, f"Missing Keywords: {', '.join(missing) or 'None'}")
    pdf.multi_cell(0, 8, f"Extra Keywords: {', '.join(extra) or 'None'}")
    pdf.multi_cell(0, 8, f"Suggestions:\n- " + "\n- ".join(suggestions)
                   if suggestions else "Suggestions: None")

    # Output PDF as a string
    pdf_bytes = pdf.output(dest="S").encode("latin1")
    return BytesIO(pdf_bytes)


# -------------------- Streamlit UI --------------------

st.set_page_config(page_title="Resume ↔ Job Matcher", layout='wide')
st.title("🧾 Resume vs Job Description Matcher")
st.markdown("Upload a resume and a job description. The app will extract keywords, compute a match score, and give action items to improve your resume.")

with st.sidebar:
    st.header("Settings")
    fuzzy_threshold = st.slider("Fuzzy match threshold", 60, 100, 82)
    use_builtin_skills = st.checkbox(
        "Use built-in skills taxonomy", value=True)
    show_extra = st.checkbox("Show extra resume keywords", value=True)
    export_pdf = st.checkbox("Enable export to PDF", value=True)

col1, col2 = st.columns(2)
with col1:
    uploaded_resume = st.file_uploader(
        "Upload Resume (PDF/DOCX/TXT)", type=['pdf', 'docx', 'txt'])
    resume_text, resume_name = load_file_to_text(uploaded_resume)
    st.text_input("Optional: Resume filename for report",
                  value=resume_name or "resume")

with col2:
    uploaded_jd = st.file_uploader(
        "Upload Job Description (PDF/DOCX/TXT)", type=['pdf', 'docx', 'txt'])
    jd_text, jd_name = load_file_to_text(uploaded_jd)
    st.text_input("Optional: Job Description filename for report",
                  value=jd_name or "job_description")

if not resume_text or not jd_text:
    st.info("Upload both a resume and a job description to get a match.")
    st.stop()

# Build job keywords
with st.spinner("Analyzing job description..."):
    job_keywords, job_freq = build_job_keywords(
        jd_text, BUILTIN_SKILLS if use_builtin_skills else set(), nlp, fuzzy_threshold=fuzzy_threshold)

with st.spinner("Analyzing resume and computing match..."):
    match_score, matched, missing, extra = compute_match(
        resume_text, job_keywords, fuzzy_threshold=fuzzy_threshold)
    suggestions = generate_suggestions(missing, extra)

# Display results
st.subheader(f"Match Score: {match_score}%")
colA, colB = st.columns([2, 3])
with colA:
    st.metric("Match", f"{match_score}%", delta=None)
    st.write("**Top suggestions**")
    for s in suggestions:
        st.write(f"- {s}")
    if export_pdf:
        pdf_bytes = export_report_pdf(
            match_score, matched, missing, extra, suggestions)
        st.download_button("Download PDF Report", data=pdf_bytes,
                           file_name="match_report.pdf", mime='application/pdf')

with colB:
    st.markdown("**Matched keywords (found in resume)**")
    if matched:
        st.write(', '.join(matched[:200]))
    else:
        st.write("None found")
    st.markdown(
        "**Missing keywords (present in job description but not in resume)**")
    if missing:
        # color-coded chips
        for kw in missing:
            st.markdown(
                f"<span style='background-color:#807f82;padding:4px;border-radius:4px;margin:2px;display:inline-block'>{kw}</span>", unsafe_allow_html=True)
    else:
        st.write("None 🎉")

    if show_extra:
        st.markdown("**Extra keywords in resume (not in job description)**")
        if extra:
            st.write(', '.join(extra[:200]))
        else:
            st.write("None")

st.markdown("---")
st.write("**How it works (short):** The app extracts text from files, uses spaCy to lemmatize and pull candidate terms, then compares job keywords (from a skills taxonomy + noun chunks) to the resume using fuzzy matching (RapidFuzz). The match is the ratio of matched job keywords to all job keywords.")

st.header("Tweak & Re-run")
st.write("If results seem off, try: lowering the fuzzy threshold, adding more entries to the built-in skills list, or manually adding key skills into the job description input before re-running.")

st.info("Want this as a CLI script or a Flask app instead? Reply and I will adapt the code.")
